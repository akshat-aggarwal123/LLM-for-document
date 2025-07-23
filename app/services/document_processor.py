"""Document processing service for extracting and chunking text from various file formats"""

import os
import re
from typing import List, Dict, Any, Tuple
from pathlib import Path
import time

# Document processing libraries
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import email

from app.core.config import DocumentClause, CLAUSE_TYPES
from app.core.exceptions import DocumentProcessingError, FileTypeNotSupportedError
from app.utils.logging import app_logger as logger


class DocumentProcessor:
    """Handles document loading, text extraction, and chunking"""

    def __init__(self, settings):
        print(f"[DEBUG] DocumentProcessor.__init__ - settings type: {type(settings)}")
        self.settings = settings
        self.supported_extensions = ['.pdf', '.docx', '.txt', '.html', '.eml']
        print(f"[DEBUG] DocumentProcessor initialized with settings: {self.settings is not None}")

    def process_document(self, file_path: str) -> List[DocumentClause]:
        """
        Process a document and return structured clauses

        Args:
            file_path: Path to the document file

        Returns:
            List of DocumentClause objects
        """
        print(f"[DEBUG] DocumentProcessor.process_document called with: {file_path}")
        start_time = time.time()

        try:
            # Validate file
            print("[DEBUG] Starting file validation...")
            self._validate_file(file_path)
            print("[DEBUG] File validation completed")

            # Extract text content
            print("[DEBUG] Starting text extraction...")
            pages_content = self._extract_text(file_path)
            print(f"[DEBUG] Text extraction completed - {len(pages_content)} pages extracted")

            # Create chunks from content
            print("[DEBUG] Starting clause creation...")
            clauses = self._create_clauses(pages_content, file_path)
            print(f"[DEBUG] Clause creation completed - {len(clauses)} clauses created")

            processing_time = time.time() - start_time
            logger.info(f"Processed {len(clauses)} clauses from {file_path} in {processing_time:.2f}s")
            print(f"[DEBUG] Total processing time: {processing_time:.2f}s")

            return clauses

        except Exception as e:
            processing_time = time.time() - start_time
            print(f"[DEBUG] Error in process_document after {processing_time:.2f}s: {str(e)}")
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")

    def _validate_file(self, file_path: str) -> None:
        """Validate file existence, type, and size"""
        print(f"[DEBUG] _validate_file called with: {file_path}")

        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            print(f"[DEBUG] {error_msg}")
            raise DocumentProcessingError(error_msg)

        file_ext = Path(file_path).suffix.lower()
        print(f"[DEBUG] File extension: {file_ext}")
        
        if file_ext not in self.supported_extensions:
            error_msg = f"Unsupported file type: {file_ext}"
            print(f"[DEBUG] {error_msg}")
            raise FileTypeNotSupportedError(error_msg)

        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        print(f"[DEBUG] File size: {file_size_mb:.2f}MB")
        
        # FIX: Use self.settings instead of settings
        max_file_size = getattr(self.settings, 'max_file_size', 100)  # Default to 100MB if not set
        print(f"[DEBUG] Max file size limit: {max_file_size}MB")
        
        if file_size_mb > max_file_size:
            error_msg = f"File size ({file_size_mb:.1f}MB) exceeds limit ({max_file_size}MB)"
            print(f"[DEBUG] {error_msg}")
            raise DocumentProcessingError(error_msg)

        print("[DEBUG] File validation passed")

    def _extract_text(self, file_path: str) -> List[Tuple[str, int]]:
        """
        Extract text content from file

        Returns:
            List of tuples (text_content, page_number)
        """
        print(f"[DEBUG] _extract_text called with: {file_path}")
        file_ext = Path(file_path).suffix.lower()
        print(f"[DEBUG] File extension for extraction: {file_ext}")

        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        elif file_ext == '.txt':
            return self._extract_from_txt(file_path)
        elif file_ext == '.html':
            return self._extract_from_html(file_path)
        elif file_ext == '.eml':
            return self._extract_from_email(file_path)
        else:
            error_msg = f"Unsupported file extension: {file_ext}"
            print(f"[DEBUG] {error_msg}")
            raise FileTypeNotSupportedError(error_msg)

    def _extract_from_pdf(self, file_path: str) -> List[Tuple[str, int]]:
        """Extract text from PDF file with improved content processing"""
        print(f"[DEBUG] _extract_from_pdf called")
        pages_content = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                print(f"[DEBUG] PDF has {len(pdf_reader.pages)} pages")

                current_section = []
                current_page = 1
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    print(f"[DEBUG] Processing PDF page {page_num}")
                    text = page.extract_text()
                    
                    if text.strip():
                        # Basic cleanup while preserving content
                        text = text.strip()
                        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
                        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)  # Remove control chars but keep unicode
                        
                        # Keep track of section boundaries
                        sections = []
                        current_text = ""
                        
                        # Split by newlines and process each line
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line:
                                continue
                                
                            # Check if line is a header/section marker
                            is_header = re.match(
                                r'(?i)^(?:section|coverage|benefit|exclusion|limitation|policy|insurance|medical|dental|vision)',
                                line
                            )
                            
                            if is_header and current_text:
                                # Save current section if it exists
                                sections.append(current_text)
                                current_text = line
                            else:
                                # Add line to current section
                                current_text = current_text + ' ' + line if current_text else line
                        
                        # Add final section
                        if current_text:
                            sections.append(current_text)
                        
                        # Save each section with page number
                        for section in sections:
                            if len(section.split()) > 5:  # Only skip very short sections
                                pages_content.append((section, page_num))
                                print(f"[DEBUG] Added section from page {page_num}: {len(section)} chars")
                    else:
                        print(f"[DEBUG] Page {page_num}: no text extracted")

                # Don't forget the last section
                if current_section:
                    section_text = ' '.join(current_section)
                    if len(section_text.split()) > 10:
                        pages_content.append((section_text, current_page))
                        print(f"[DEBUG] Added final section from page {current_page}: {len(section_text)} chars")

        except Exception as e:
            error_msg = f"Failed to extract text from PDF: {str(e)}"
            print(f"[DEBUG] PDF extraction error: {error_msg}")
            raise DocumentProcessingError(error_msg)

        print(f"[DEBUG] PDF extraction completed: {len(pages_content)} pages with content")
        return pages_content

    def _extract_from_docx(self, file_path: str) -> List[Tuple[str, int]]:
        """Extract text from DOCX file"""
        print(f"[DEBUG] _extract_from_docx called")
        try:
            doc = Document(file_path)
            text_content = []

            print(f"[DEBUG] DOCX has {len(doc.paragraphs)} paragraphs")
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    print(f"[DEBUG] Paragraph {i+1}: {len(paragraph.text)} characters")

            # Combine all text into a single page (DOCX doesn't have clear page breaks)
            full_text = '\n'.join(text_content)
            result = [(full_text, 1)] if full_text else []
            print(f"[DEBUG] DOCX extraction completed: {len(full_text) if full_text else 0} characters")
            return result

        except Exception as e:
            error_msg = f"Failed to extract text from DOCX: {str(e)}"
            print(f"[DEBUG] DOCX extraction error: {error_msg}")
            raise DocumentProcessingError(error_msg)

    def _extract_from_txt(self, file_path: str) -> List[Tuple[str, int]]:
        """Extract text from TXT file"""
        print(f"[DEBUG] _extract_from_txt called")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
                result = [(content, 1)] if content else []
                print(f"[DEBUG] TXT extraction completed: {len(content) if content else 0} characters")
                return result

        except Exception as e:
            error_msg = f"Failed to extract text from TXT: {str(e)}"
            print(f"[DEBUG] TXT extraction error: {error_msg}")
            raise DocumentProcessingError(error_msg)

    def _extract_from_html(self, file_path: str) -> List[Tuple[str, int]]:
        """Extract text from HTML file"""
        print(f"[DEBUG] _extract_from_html called")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                text = soup.get_text()
                result = [(text.strip(), 1)] if text.strip() else []
                print(f"[DEBUG] HTML extraction completed: {len(text) if text else 0} characters")
                return result

        except Exception as e:
            error_msg = f"Failed to extract text from HTML: {str(e)}"
            print(f"[DEBUG] HTML extraction error: {error_msg}")
            raise DocumentProcessingError(error_msg)

    def _extract_from_email(self, file_path: str) -> List[Tuple[str, int]]:
        """Extract text from email file"""
        print(f"[DEBUG] _extract_from_email called")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                msg = email.message_from_file(file)

            # Extract subject and body
            subject = msg.get('Subject', '')
            print(f"[DEBUG] Email subject: {subject[:50]}...")

            body = ""
            if msg.is_multipart():
                print("[DEBUG] Processing multipart email")
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body += part.get_payload(decode=True).decode('utf-8')
            else:
                print("[DEBUG] Processing simple email")
                body = msg.get_payload(decode=True).decode('utf-8')

            content = f"Subject: {subject}\n\n{body}".strip()
            result = [(content, 1)] if content else []
            print(f"[DEBUG] Email extraction completed: {len(content) if content else 0} characters")
            return result

        except Exception as e:
            error_msg = f"Failed to extract text from email: {str(e)}"
            print(f"[DEBUG] Email extraction error: {error_msg}")
            raise DocumentProcessingError(error_msg)

    def _create_clauses(self, pages_content: List[Tuple[str, int]], file_path: str) -> List[DocumentClause]:
        """
        Create DocumentClause objects from extracted text

        Args:
            pages_content: List of (text, page_number) tuples
            file_path: Source file path

        Returns:
            List of DocumentClause objects
        """
        print(f"[DEBUG] _create_clauses called with {len(pages_content)} pages")
        clauses = []
        clause_id = 0

        for page_idx, (text_content, page_num) in enumerate(pages_content):
            print(f"[DEBUG] Processing page {page_idx + 1} (page number {page_num})")
            
            # Split text into chunks
            chunks = self._chunk_text(text_content)
            print(f"[DEBUG] Page {page_num}: created {len(chunks)} chunks")

            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk.strip()) < 20:  # Skip very short chunks
                    print(f"[DEBUG] Skipping short chunk {chunk_idx + 1} (length: {len(chunk.strip())})")
                    continue

                clause_id += 1
                clause_type = self._classify_clause(chunk)
                print(f"[DEBUG] Clause {clause_id}: type={clause_type}, length={len(chunk)}")

                clause = DocumentClause(
                    id=f"{Path(file_path).stem}_{clause_id}",
                    text=chunk,
                    clause_type=clause_type,
                    source_document=os.path.basename(file_path),
                    page_number=page_num,
                    metadata={
                        'chunk_length': len(chunk),
                        'file_path': file_path
                    }
                )
                clauses.append(clause)

        print(f"[DEBUG] Created {len(clauses)} total clauses")
        return clauses

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks using semantic boundaries and preserving insurance-related context

        Args:
            text: Input text to chunk

        Returns:
            List of text chunks
        """
        print(f"[DEBUG] _chunk_text called with text length: {len(text)}")
        
        # Define insurance-specific section markers with flags moved outside
        section_markers = [
            r'what (?:does|is) .*?(?:covered|coverage|benefit).*?\?',
            r'how do I (?:know|find out|check).*?\?',
            r'(?:when|where) (?:to|should I) call.*?\?',
            r'(?:what|how) (?:type|kind) of (?:plan|coverage|insurance).*?\?'
        ]
        # Compile patterns with flags
        section_markers = [re.compile(pattern, re.IGNORECASE) for pattern in section_markers]
        
        # Split into semantic sections first
        sections = [text]
        for marker in section_markers:
            new_sections = []
            for section in sections:
                # Use the compiled pattern
                splits = marker.split(section)
                for i, split in enumerate(splits):
                    if i > 0 and marker.match(splits[i-1]):
                        # Keep the question with its answer
                        new_sections.append(splits[i-1] + split)
                    elif not marker.match(split):
                        new_sections.append(split)
            sections = new_sections

        # Then split remaining text into sentences
        chunks = []
        max_chunk_size = getattr(self.settings, 'max_chunk_size', 1000)  # Default to 1000 chars
        print(f"[DEBUG] Max chunk size: {max_chunk_size}")

        for section in sections:
            if not section.strip():
                continue
                
            current_chunk = ""
            # Split into sentences using proper regex
            # Create a proper sentence splitting pattern
            sentence_pattern = re.compile(r'([.!?])\s+')
            sentences = [s.strip() for s in sentence_pattern.split(section) if s.strip()]

            for sentence in sentences:
                if len(current_chunk) + len(sentence) < max_chunk_size:
                    current_chunk += " " + sentence if current_chunk else sentence
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        print(f"[DEBUG] Created chunk {len(chunks)}: {len(current_chunk.strip())} characters")
                    current_chunk = sentence

        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
            print(f"[DEBUG] Created final chunk {len(chunks)}: {len(current_chunk.strip())} characters")

        print(f"[DEBUG] Text chunking completed: {len(chunks)} chunks created")
        return chunks

    def _classify_clause(self, text: str) -> str:
        """
        Classify clause type based on keywords and insurance-specific patterns

        Args:
            text: Clause text

        Returns:
            Clause type string
        """
        text_lower = text.lower()

        # Insurance-specific patterns
        insurance_patterns = {
            'coverage_definition': ['covered', 'means', 'defined as'],
            'coverage_info': ['what is covered', 'services covered', 'benefits include'],
            'benefits': ['summary of benefits', 'member handbook', 'policy handbook'],
            'contact_info': ['call customer', 'member service', 'representative'],
            'plan_types': ['hmo', 'ppo', 'insurance plan', 'network provider']
        }

        # First check for insurance-specific patterns
        for clause_type, patterns in insurance_patterns.items():
            if any(pattern in text_lower for pattern in patterns):
                print(f"[DEBUG] Classified insurance clause as '{clause_type}'")
                return clause_type

        # Fall back to general classification
        scores = {}
        for clause_type, keywords in CLAUSE_TYPES.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            scores[clause_type] = score

        # Return the type with highest score, default to 'general'
        if max(scores.values()) > 0:
            best_type = max(scores, key=scores.get)
            print(f"[DEBUG] Classified clause as '{best_type}' (score: {scores[best_type]})")
            return best_type
        else:
            print(f"[DEBUG] Classified clause as 'general' (no keyword matches)")
            return 'general'

    def get_supported_file_types(self) -> List[str]:
        """Return list of supported file extensions"""
        return self.supported_extensions